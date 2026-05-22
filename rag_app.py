import os
import json
import hashlib
from datetime import datetime
from pathlib import Path  # used for _SYSTEM_FONTS font paths
import streamlit as st
import voyageai
from langchain_anthropic import ChatAnthropic
from langchain_community.document_loaders import PyPDFLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_chroma import Chroma
from langchain_core.documents import Document
from langchain_core.embeddings import Embeddings
from langchain_core.prompts import PromptTemplate
from langchain_core.callbacks import BaseCallbackHandler
from dotenv import load_dotenv
import tempfile

load_dotenv()

PERSONA_NAMES = ["Compliance Officer", "RegTech Product Manager", "Policy & Regulatory Affairs Analyst"]


class VoyageEmbeddings(Embeddings):
    def __init__(self, model: str = "voyage-3", api_key: str | None = None):
        self.client = voyageai.Client(api_key=api_key or os.environ["VOYAGE_API_KEY"])
        self.model = model

    def embed_documents(self, texts: list[str]) -> list[list[float]]:
        result = self.client.embed(texts, model=self.model, input_type="document")
        return result.embeddings

    def embed_query(self, text: str) -> list[float]:
        result = self.client.embed([text], model=self.model, input_type="query")
        return result.embeddings[0]


class StreamHandler(BaseCallbackHandler):
    def __init__(self, container):
        self.container = container
        self.text = ""

    def on_llm_new_token(self, token: str, **_):
        self.text += token
        self.container.markdown(self.text + "▌")

    def on_llm_end(self, *_, **__):
        self.container.markdown(self.text)


def analyze_document(text: str) -> dict:
    """Single Haiku call: checks compliance relevance and extracts document profile."""
    llm = ChatAnthropic(
        model="claude-haiku-4-5-20251001",
        api_key=os.environ["ANTHROPIC_API_KEY"],
        max_tokens=250,
    )
    response = llm.invoke(
        "Analyze this regulatory document excerpt. Return a JSON object with exactly these fields:\n"
        '- "is_compliance": true if the document is related to AML, financial crime compliance, '
        'or financial regulation; false otherwise\n'
        '- "jurisdictions": array of applicable jurisdictions detected from the content. '
        'Use values from: "US", "UK", "EU", "India", "Japan", "Global". '
        'Use "Global" for multi-jurisdictional documents like FATF guidance. '
        'A document may have multiple values.\n'
        '- "institution_name": the primary regulated institution, bank, or entity named in the document. '
        'Use the exact name from the text. null if not identifiable\n'
        '- "document_title": the formal title of the document. null if not identifiable\n'
        '- "document_type": short label such as "consent order", "guidance", "circular", "regulation". '
        'null if unclear\n'
        '- "reason": one short sentence explaining your determination\n\n'
        "Return only valid JSON, no markdown or extra text.\n\n"
        f"Document excerpt:\n{text[:3000]}"
    )
    try:
        content = response.content.strip().lstrip("```json").lstrip("```").rstrip("```").strip()
        return json.loads(content)
    except Exception:
        is_comp = "compliance" in response.content.lower() or "aml" in response.content.lower()
        return {
            "is_compliance": is_comp,
            "jurisdictions": ["Unknown"],
            "institution_name": None,
            "document_title": None,
            "document_type": None,
            "reason": response.content[:120],
        }


def load_pdf_pages(path: str) -> list[Document]:
    """Load PDF pages and prepend embedded PDF title metadata when available."""
    docs = PyPDFLoader(path, extraction_mode="plain").load()
    try:
        from pypdf import PdfReader

        reader = PdfReader(path)
        pdf_meta = reader.metadata or {}
        pdf_title = (pdf_meta.get("/Title") or pdf_meta.get("title") or "").strip()
        if pdf_title and docs:
            docs[0].page_content = f"{pdf_title}\n{docs[0].page_content}"
    except Exception:
        pass
    return docs


def build_document_chunks(
    pages: list[Document],
    filename: str,
    jurisdictions: list[str],
    profile: dict,
) -> list[Document]:
    """Build a searchable header chunk plus body chunks for one PDF."""
    jur_str = ", ".join(jurisdictions)
    institution = profile.get("institution_name") or "Not identified"
    title = profile.get("document_title") or filename
    doc_type = profile.get("document_type") or "Regulatory document"
    opening = "\n\n".join(page.page_content for page in pages[:3] if page.page_content.strip())

    header = Document(
        page_content=(
            f"Document file: {filename}\n"
            f"Jurisdiction: {jur_str}\n"
            f"Institution: {institution}\n"
            f"Document title: {title}\n"
            f"Document type: {doc_type}\n\n"
            f"--- Document opening ---\n{opening[:4500]}"
        ),
        metadata={
            "source": filename,
            "jurisdiction": jur_str,
            "chunk_type": "header",
            "page": 0,
            "institution_name": institution,
        },
    )

    for page in pages:
        page.metadata["source"] = filename
        page.metadata["jurisdiction"] = jur_str
        page.metadata["chunk_type"] = "body"
        page.metadata["institution_name"] = institution

    body_chunks = RecursiveCharacterTextSplitter(
        chunk_size=800,
        chunk_overlap=160,
        separators=["\n\n", "\n", ". ", " ", ""],
    ).split_documents(pages)

    return [header, *body_chunks]


def condense_question(question: str, chat_history: list[tuple[str, str]], llm) -> str:
    if not chat_history:
        return question
    history_text = "\n".join(f"Q: {q}\nA: {a}" for q, a in chat_history[-3:])
    response = llm.invoke(
        "Rewrite the follow-up question as a standalone search query using the chat history only "
        "when needed for context. Return only the rewritten question.\n\n"
        f"Chat history:\n{history_text}\n\nFollow-up: {question}\n\nStandalone question:"
    )
    rewritten = response.content.strip()
    return rewritten or question


def get_header_documents(vectorstore) -> list[Document]:
    try:
        records = vectorstore.get(where={"chunk_type": "header"})
    except Exception:
        return []
    if not records or not records.get("documents"):
        return []
    return [
        Document(page_content=content, metadata=meta or {})
        for content, meta in zip(records["documents"], records["metadatas"])
    ]


def retrieve_documents(vectorstore, question: str, k: int = 10) -> list[Document]:
    header_docs = get_header_documents(vectorstore)
    body_docs = vectorstore.similarity_search(question, k=k)

    seen: set[str] = set()
    merged: list[Document] = []
    for doc in header_docs + body_docs:
        key = doc.page_content[:240]
        if key in seen:
            continue
        seen.add(key)
        merged.append(doc)
    return merged


def format_context(docs: list[Document]) -> str:
    parts = []
    for doc in docs:
        source = doc.metadata.get("source", "Unknown")
        page = doc.metadata.get("page", "?")
        chunk_type = doc.metadata.get("chunk_type", "body")
        label = "Document summary / opening" if chunk_type == "header" else f"page {page}"
        parts.append(f"[{source} | {label}]\n{doc.page_content}")
    return "\n\n---\n\n".join(parts)


def evaluate_answer_confidence(question: str, context: str, answer: str, llm) -> dict:
    """Ask the model how well the answer is supported by the retrieved source excerpts."""
    response = llm.invoke(
        "Evaluate whether the assistant's answer is correct and grounded in the source excerpts. "
        "Judge only from the excerpts — do not use outside knowledge.\n\n"
        "Return a JSON object with exactly these fields:\n"
        '- "confidence_score": number from 0.0 to 1.0 — how confident you are the answer is right '
        "and supported by the sources\n"
        '- "confidence_level": "high", "medium", or "low"\n'
        "  - high (>=0.75): key claims are explicitly supported by the excerpts\n"
        "  - medium (0.5-0.74): partially supported, relies on inference, or excerpts are incomplete\n"
        "  - low (<0.5): weak support, missing evidence, contradictions, or likely hallucination\n"
        '- "reason": 1-2 sentences explaining why you chose this score — say what supported the answer '
        "or what was missing/contradicted\n\n"
        "Return only valid JSON, no markdown.\n\n"
        f"Question:\n{question}\n\n"
        f"Source excerpts:\n{context[:8000]}\n\n"
        f"Answer:\n{answer}"
    )
    try:
        content = response.content.strip().lstrip("```json").lstrip("```").rstrip("```").strip()
        result = json.loads(content)
        score = max(0.0, min(1.0, float(result.get("confidence_score", 0.5))))
        level = str(result.get("confidence_level", "")).lower()
        if level not in {"high", "medium", "low"}:
            level = "high" if score >= 0.75 else "medium" if score >= 0.5 else "low"
        reason = str(result.get("reason", "")).strip() or "No explanation provided."
        return {
            "confidence_score": score,
            "confidence_level": level,
            "confidence_msg": reason,
        }
    except Exception:
        return {
            "confidence_score": 0.5,
            "confidence_level": "medium",
            "confidence_msg": "Could not evaluate answer confidence automatically.",
        }


def export_chat_as_markdown(
    messages: list,
    persona: str,
    jurisdictions: list[str],
    indexed_docs: list,
) -> str:
    lines = [
        "# AML Regulatory Knowledge Base — Chat Export",
        "",
        f"**Exported:** {datetime.now().strftime('%Y-%m-%d %H:%M')}  ",
        f"**Persona:** {persona}  ",
        f"**Jurisdictions:** {', '.join(jurisdictions)}  ",
        "",
        "## Indexed Documents",
    ]
    for doc in indexed_docs:
        badges = " ".join(f"`{j}`" for j in doc.get("jurisdictions", []))
        lines.append(f"- {doc['name']} {badges}")
    lines += ["", "---", "", "## Conversation", ""]

    for msg in messages:
        if msg["role"] == "user":
            lines += [f"**You:** {msg['content']}", ""]
        else:
            lines += [f"**Assistant:** {msg['content']}", ""]
            level = msg.get("confidence_level", "")
            if level:
                score_pct = round(msg.get("confidence_score", 0) * 100)
                conf_line = f"Confidence: {score_pct}% · {level.capitalize()}"
                if msg.get("confidence_msg"):
                    conf_line += f" — {msg['confidence_msg']}"
                lines += [f"> _{conf_line}_", ""]
            if msg.get("sources"):
                lines.append("*Sources:*")
                for src in msg["sources"]:
                    lines.append(f"- {src}")
            lines += ["", "---", ""]

    return "\n".join(lines)


def export_chat_as_docx(messages: list, persona: str, jurisdictions: list[str], indexed_docs: list) -> bytes:
    from docx import Document
    from docx.shared import RGBColor
    from io import BytesIO

    doc = Document()
    doc.add_heading("AML Regulatory Knowledge Base — Chat Export", 0)

    meta = doc.add_paragraph()
    meta.add_run("Exported: ").bold = True
    meta.add_run(datetime.now().strftime("%Y-%m-%d %H:%M"))
    meta = doc.add_paragraph()
    meta.add_run("Persona: ").bold = True
    meta.add_run(persona)
    meta = doc.add_paragraph()
    meta.add_run("Jurisdictions: ").bold = True
    meta.add_run(", ".join(jurisdictions))

    doc.add_heading("Indexed Documents", 1)
    for d in indexed_docs:
        badges = ", ".join(d.get("jurisdictions", []))
        doc.add_paragraph(f"{d['name']}  [{badges}]", style="List Bullet")

    doc.add_heading("Conversation", 1)
    for msg in messages:
        if msg["role"] == "user":
            p = doc.add_paragraph()
            p.add_run("You: ").bold = True
            p.add_run(msg["content"])
        else:
            p = doc.add_paragraph()
            p.add_run("Assistant: ").bold = True
            p.add_run(msg["content"])
            if msg.get("confidence_level"):
                score_pct = round(msg.get("confidence_score", 0) * 100)
                level = msg["confidence_level"].capitalize()
                conf_text = f"Confidence: {score_pct}% · {level}"
                if msg.get("confidence_msg"):
                    conf_text += f" — {msg['confidence_msg']}"
                cp = doc.add_paragraph(conf_text)
                cp.runs[0].italic = True
                cp.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)
            if msg.get("sources"):
                sp = doc.add_paragraph()
                sp.add_run("Sources: ").italic = True
                for src in msg["sources"]:
                    doc.add_paragraph(src, style="List Bullet")
        doc.add_paragraph()

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


_SYSTEM_FONTS = {
    "": Path("C:/Windows/Fonts/arial.ttf"),
    "B": Path("C:/Windows/Fonts/arialbd.ttf"),
    "I": Path("C:/Windows/Fonts/ariali.ttf"),
}


def export_chat_as_pdf(messages: list, persona: str, jurisdictions: list[str], indexed_docs: list) -> bytes:
    from fpdf import FPDF

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    for style, path in _SYSTEM_FONTS.items():
        pdf.add_font("Arial", style=style, fname=str(path))

    pdf.add_page()

    pdf.set_font("Arial", "B", 16)
    pdf.multi_cell(0, 10, "AML Regulatory Knowledge Base — Chat Export")
    pdf.ln(2)
    pdf.set_font("Arial", "", 10)
    pdf.cell(0, 6, f"Exported: {datetime.now().strftime('%Y-%m-%d %H:%M')}", ln=True)
    pdf.cell(0, 6, f"Persona: {persona}", ln=True)
    pdf.cell(0, 6, f"Jurisdictions: {', '.join(jurisdictions)}", ln=True)
    pdf.ln(4)

    pdf.set_font("Arial", "B", 12)
    pdf.cell(0, 8, "Indexed Documents", ln=True)
    pdf.set_font("Arial", "", 10)
    for d in indexed_docs:
        badges = ", ".join(d.get("jurisdictions", []))
        pdf.cell(0, 6, f"  - {d['name']}  [{badges}]", ln=True)
    pdf.ln(4)

    pdf.set_font("Arial", "B", 12)
    pdf.cell(0, 8, "Conversation", ln=True)
    pdf.ln(2)

    for msg in messages:
        if msg["role"] == "user":
            pdf.set_font("Arial", "B", 10)
            pdf.multi_cell(0, 6, f"You:  {msg['content']}")
        else:
            pdf.set_font("Arial", "B", 10)
            pdf.cell(0, 6, "Assistant:", ln=True)
            pdf.set_font("Arial", "", 10)
            pdf.multi_cell(0, 6, msg["content"])
            if msg.get("confidence_level"):
                score_pct = round(msg.get("confidence_score", 0) * 100)
                level = msg["confidence_level"].capitalize()
                conf_text = f"Confidence: {score_pct}% - {level}"
                if msg.get("confidence_msg"):
                    conf_text += f" - {msg['confidence_msg']}"
                pdf.set_font("Arial", "I", 9)
                pdf.multi_cell(0, 5, conf_text)
            if msg.get("sources"):
                pdf.set_font("Arial", "I", 9)
                pdf.cell(0, 5, "Sources:", ln=True)
                for src in msg["sources"]:
                    pdf.multi_cell(0, 5, f"  - {src}")
        pdf.set_draw_color(200, 200, 200)
        pdf.line(pdf.l_margin, pdf.get_y() + 3, pdf.w - pdf.r_margin, pdf.get_y() + 3)
        pdf.ln(7)

    return bytes(pdf.output())


def get_persona_prompt(persona: str, jurisdictions: list[str]) -> PromptTemplate:
    jur = ", ".join(jurisdictions) if jurisdictions else "the available jurisdictions"
    grounding = (
        "Answer ONLY using the Context from regulatory documents below. "
        "Each context block is labeled with the source document and page or section. "
        "For factual questions about names, titles, dates, or parties, quote directly from the "
        "Document summary / opening sections when present. "
        "If the context does not contain enough information, say so explicitly. "
        "Do not use outside knowledge or reference regulators, rules, or jurisdictions "
        "not present in the context.\n\n"
    )
    templates = {
        "Compliance Officer": f"""You are an AML compliance expert advising a compliance officer at a financial institution.
The knowledge base covers the following jurisdictions: {jur}.
Focus on practical obligations, examination risk, and what regulators expect in practice.
Be direct. Cite specific thresholds, timeframes, and trigger criteria where they exist.
Flag where non-compliance creates material regulatory or reputational risk.
For cross-jurisdictional questions, explicitly compare and contrast requirements.

{grounding}Context from regulatory documents:
{{context}}

Question: {{question}}
Answer:""",

        "RegTech Product Manager": f"""You are an AML regulatory expert advising a product manager building compliance systems.
The knowledge base covers the following jurisdictions: {jur}.
Translate regulatory requirements into product specifications — data fields, workflow logic, thresholds, audit trail requirements.
Structure answers as: requirement → what to build → edge cases to handle.
Avoid regulatory jargon; use systems and product language.
For cross-jurisdictional questions, highlight where requirements diverge and where a single implementation can satisfy multiple regimes.

{grounding}Context from regulatory documents:
{{context}}

Question: {{question}}
Answer:""",

        "Policy & Regulatory Affairs Analyst": f"""You are an AML regulatory expert advising a policy and regulatory affairs analyst.
The knowledge base covers the following jurisdictions: {jur}.
Focus on regulatory intent, jurisdictional comparisons, interpretive nuance, and gaps in the framework.
Use precise regulatory language. Note where rules are ambiguous or where guidance diverges from legislation.
Reference cross-jurisdictional parallels and cite section numbers where relevant.

{grounding}Context from regulatory documents:
{{context}}

Question: {{question}}
Answer:""",
    }
    return PromptTemplate(input_variables=["context", "question"], template=templates[persona])


def clear_knowledge_base() -> None:
    for key in ("vectorstore", "indexed_hash", "indexed_docs", "messages", "chat_history", "chat_index_hash"):
        st.session_state.pop(key, None)


st.set_page_config(page_title="AML Knowledge Base", page_icon="⚖️", layout="wide")

# --- Sidebar ---
with st.sidebar:
    st.header("Documents")
    st.caption(
        "Upload English-language PDFs. Jurisdiction is auto-detected. "
        "Remove old files from the list before uploading a new set — all selected PDFs form your active knowledge base."
    )

    if st.session_state.get("vectorstore") and st.button("Reset knowledge base", use_container_width=True):
        clear_knowledge_base()
        st.rerun()

    uploaded_files = st.file_uploader(
        "Upload PDFs",
        type="pdf",
        accept_multiple_files=True,
        label_visibility="collapsed",
    )

    if uploaded_files:
        content_hash = hashlib.md5()
        for f in sorted(uploaded_files, key=lambda x: x.name):
            content_hash.update(f.name.encode())
            content_hash.update(f.getvalue())
        file_hash = content_hash.hexdigest()

        if st.session_state.get("indexed_hash") != file_hash:
            all_docs = []
            accepted = []

            with st.status("Processing documents...", expanded=True) as status:
                for uf in uploaded_files:
                    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
                        tmp.write(uf.getvalue())
                        tmp_path = tmp.name
                    docs = load_pdf_pages(tmp_path)

                    st.write(f"Analysing **{uf.name}**...")
                    sample = "\n\n".join(d.page_content for d in docs[:3])
                    result = analyze_document(sample)

                    if not result["is_compliance"]:
                        reason = result.get("reason", "")
                        st.warning(f"**{uf.name}** skipped — {reason}")
                        continue

                    jurisdictions = result.get("jurisdictions") or ["Unknown"]
                    jur_str = ", ".join(jurisdictions)
                    doc_chunks = build_document_chunks(docs, uf.name, jurisdictions, result)
                    all_docs.extend(doc_chunks)
                    accepted.append({
                        "name": uf.name,
                        "jurisdictions": jurisdictions,
                        "institution_name": result.get("institution_name"),
                        "document_title": result.get("document_title"),
                    })
                    institution = result.get("institution_name") or "Not identified"
                    st.write(
                        f"✓ **{uf.name}** — `{jur_str}` · "
                        f"Institution: **{institution}**"
                    )

                if not all_docs:
                    status.update(label="No compliance documents found.", state="error")
                    clear_knowledge_base()
                else:
                    embeddings = VoyageEmbeddings(model="voyage-3")
                    st.session_state.vectorstore = Chroma.from_documents(
                        all_docs,
                        embeddings,
                        collection_name=f"kb_{file_hash[:16]}",
                    )
                    st.session_state.indexed_hash = file_hash
                    st.session_state.indexed_docs = accepted
                    st.session_state.messages = []
                    st.session_state.chat_history = []
                    st.session_state.chat_index_hash = file_hash
                    status.update(
                        label=f"Ready — {len(all_docs)} chunks from {len(accepted)} document(s). Chat cleared.",
                        state="complete",
                    )
    elif st.session_state.get("vectorstore"):
        clear_knowledge_base()
        st.rerun()

    if st.session_state.get("indexed_docs"):
        st.markdown("**Indexed documents:**")
        for doc in st.session_state.indexed_docs:
            badges = " ".join(f"`{j}`" for j in doc["jurisdictions"])
            institution = doc.get("institution_name")
            institution_label = f" · {institution}" if institution else ""
            st.markdown(f"- {doc['name']} &nbsp; {badges}{institution_label}")

    st.divider()

    st.header("Persona")
    st.caption("Tailors answer style and framing to your role.")
    st.selectbox("I am a...", PERSONA_NAMES, key="persona")

# --- Main area ---
st.title("AML Regulatory Knowledge Base")

has_index = "vectorstore" in st.session_state
with st.expander("How to use this tool", expanded=not has_index):
    st.markdown(
        "Query your AML and financial crime regulatory PDFs in a multi-turn chat. "
        "Upload documents from multiple jurisdictions and ask cross-border questions in one session."
    )
    st.markdown("**Steps**")
    st.markdown(
        "1. **Upload PDFs** in the sidebar — English-language AML/regulatory documents only\n"
        "2. **Select a persona** to tailor how answers are framed\n"
        "3. **Ask questions** across all indexed documents; follow-ups use chat history\n"
        "4. **Export the conversation** as Markdown or Word when you're done"
    )
    st.markdown("**Personas**")
    st.markdown(
        "- **Compliance Officer** — practical obligations, thresholds, and examination risk\n"
        "- **RegTech Product Manager** — product specs, data fields, workflow logic, edge cases\n"
        "- **Policy & Regulatory Affairs Analyst** — regulatory intent, comparisons, interpretive nuance"
    )
    st.markdown("**Good to know**")
    st.markdown(
        "- Each document is screened for AML relevance before indexing; unrelated files are skipped\n"
        "- Jurisdiction (US, UK, EU, India, Japan, Global) is auto-detected from document content\n"
        "- Remove old PDFs from the uploader before loading a new set — all selected files form one knowledge base\n"
        "- Changing documents clears chat history automatically\n"
        "- Answers include a confidence score with an explanation of how well they are supported by your documents\n"
        "- Session data resets on page refresh — export your chat if you need to keep it"
    )

if "messages" not in st.session_state:
    st.session_state.messages = []
if "chat_history" not in st.session_state:
    st.session_state.chat_history = []

if "vectorstore" not in st.session_state:
    st.info("Upload documents in the sidebar to get started.")
else:
    active_jurisdictions = list({
        j
        for doc in st.session_state.get("indexed_docs", [])
        for j in doc.get("jurisdictions", [])
    })

    col_title, col_export, col_clear = st.columns([3, 1.2, 1.2])
    with col_title:
        st.subheader("Chat")
    with col_export:
        has_messages = bool(st.session_state.messages)
        with st.popover("Export chat", disabled=not has_messages, use_container_width=True):
            fmt = st.radio(
                "Format",
                ["Markdown (.md)", "Word (.docx)"],
                label_visibility="collapsed",
            )
            ts = datetime.now().strftime("%Y%m%d_%H%M")
            args = (
                st.session_state.messages,
                st.session_state.get("persona", PERSONA_NAMES[0]),
                active_jurisdictions,
                st.session_state.get("indexed_docs", []),
            )
            if fmt == "Markdown (.md)":
                data, fname, mime = export_chat_as_markdown(*args), f"aml_chat_{ts}.md", "text/markdown"
            else:
                data = export_chat_as_docx(*args)
                fname = f"aml_chat_{ts}.docx"
                mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            st.download_button("Download", data=data, file_name=fname, mime=mime, use_container_width=True)
    with col_clear:
        if st.button("Clear chat", use_container_width=True):
            st.session_state.messages = []
            st.session_state.chat_history = []
            st.session_state.chat_index_hash = st.session_state.get("indexed_hash")
            st.rerun()

    if not st.session_state.messages:
        jur_str = ", ".join(active_jurisdictions)
        persona = st.session_state.get("persona", PERSONA_NAMES[0])
        st.markdown(
            f"Knowledge base ready — **{jur_str}** documents indexed. "
            f"Responding as **{persona}**. Ask me anything."
        )

    for msg in st.session_state.messages:
        with st.chat_message(msg["role"]):
            st.markdown(msg["content"])
            if msg.get("confidence_level"):
                score_pct = round(msg.get("confidence_score", 0) * 100)
                level = msg["confidence_level"].capitalize()
                conf_text = f"Confidence: {score_pct}% · {level}"
                if msg.get("confidence_msg"):
                    conf_text += f" — {msg['confidence_msg']}"
                if msg["confidence_level"] == "low":
                    st.warning(conf_text)
                elif msg["confidence_level"] == "medium":
                    st.caption(conf_text)
                else:
                    st.caption(conf_text)
            if msg.get("sources"):
                with st.expander("Sources"):
                    for src in msg["sources"]:
                        st.caption(src)

    if question := st.chat_input("Ask a question about your documents..."):
        if st.session_state.get("chat_index_hash") != st.session_state.get("indexed_hash"):
            st.session_state.messages = []
            st.session_state.chat_history = []
            st.session_state.chat_index_hash = st.session_state.get("indexed_hash")

        st.session_state.messages.append({"role": "user", "content": question})
        with st.chat_message("user"):
            st.markdown(question)

        with st.chat_message("assistant"):
            placeholder = st.empty()
            stream_handler = StreamHandler(placeholder)

            llm = ChatAnthropic(
                model="claude-sonnet-4-6",
                api_key=os.environ["ANTHROPIC_API_KEY"],
                streaming=True,
                callbacks=[stream_handler],
            )
            llm_condense = ChatAnthropic(
                model="claude-haiku-4-5-20251001",
                api_key=os.environ["ANTHROPIC_API_KEY"],
            )

            persona = st.session_state.get("persona", PERSONA_NAMES[0])
            prompt = get_persona_prompt(persona, active_jurisdictions)
            search_query = condense_question(
                question,
                st.session_state.chat_history,
                llm_condense,
            )
            source_documents = retrieve_documents(
                st.session_state.vectorstore,
                search_query,
                k=10,
            )
            context = format_context(source_documents)
            llm.invoke(prompt.format(context=context, question=question))
            answer = stream_handler.text

            sources = []
            for doc in source_documents:
                if doc.metadata.get("chunk_type") == "header":
                    location = "Document opening"
                else:
                    location = f"Page {doc.metadata.get('page', '?')}"
                sources.append(
                    f"{doc.metadata.get('source', 'Unknown')} — {location} "
                    f"[{doc.metadata.get('jurisdiction', '?')}]"
                )
            sources = list(dict.fromkeys(sources))

            with st.spinner("Checking answer against sources..."):
                evaluation = evaluate_answer_confidence(question, context, answer, llm_condense)

            conf_level = evaluation["confidence_level"]
            conf_msg = evaluation["confidence_msg"]
            avg_score = evaluation["confidence_score"]

            score_pct = round(avg_score * 100)
            level_display = conf_level.capitalize()
            conf_text = f"Confidence: {score_pct}% · {level_display} — {conf_msg}"
            if conf_level == "low":
                st.warning(conf_text)
            elif conf_level == "medium":
                st.caption(conf_text)
            else:
                st.caption(conf_text)

            if sources:
                with st.expander("Sources"):
                    for src in sources:
                        st.caption(src)

        st.session_state.messages.append({
            "role": "assistant",
            "content": answer,
            "sources": sources,
            "confidence_score": avg_score,
            "confidence_level": conf_level,
            "confidence_msg": conf_msg,
        })
        st.session_state.chat_history.append((question, answer))
        st.session_state.chat_index_hash = st.session_state.get("indexed_hash")
